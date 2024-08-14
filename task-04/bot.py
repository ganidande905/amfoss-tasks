import os
import telebot
import io
import csv
from dotenv import load_dotenv
import requests
from telebot import types
from docx import Document

load_dotenv()

BOT_TOKEN = os.environ.get('BOT_TOKEN')
GOOGLE_BOOKS_API_KEY = os.environ.get('GOOGLE_BOOKS_API_KEY')

bot =telebot.TeleBot(BOT_TOKEN)

def search_books(query):
    url = f"https://www.googleapis.com/books/v1/volumes?q={query}&key={GOOGLE_BOOKS_API_KEY}"
    try:
        response = requests.get(url)
        response.raise_for_status()
        return response.json()
    except requests.RequestException as e:
        print(f"Error fetching book data: {e}")
        return {}
@bot.message_handler(commands=['start'])
def getting_started(message):
    welcome_message = (
        "Hey hey! Welcome to PagePal 📚✨\n\n"
        "I'm your bookish sidekick! Tap /help to see all the cool stuff we can do together."
    
    )
    bot.reply_to(message,welcome_message)

@bot.message_handler(commands=['help'])
def help_commands(message):
    help_message =(
        "Here’s the 411 on what I can do for you:\n\n"
        "/start - Kick things off with a warm welcome!\n"
        "/book <genre> - Need some lit recommendations? Tell me your fave genre and I'll drop some fire book picks.\n"
        "/preview <book name> - Wanna peek at a book before committing? I got you with a preview link!\n"
        "/list - Manage your reading list like a pro.\n"
        "/help - Need a hand? You’re already in the right place!.\n"
        "/bye - If you gonna leave me in between , I least deserve a bye!"
    )
    bot.reply_to(message,help_message)

@bot.message_handler(commands=['bye'])
def bye_message(message):
    bye = ("Catch you later, fam! 👋✨ Stay awesome!")
    bot.reply_to(message,bye)

@bot.message_handler(commands=['book'])
def book_recommend(message):
    genre = message.text[len('\book '):].strip()
    if not genre:
        bot.reply_to(message,"Oops! You forgot to drop a genre. Try something like /book sci-fi to get some cool book recs.")
        return
    results = search_books(genre)
    books = results.get('items',[])
    if not books:
        bot.reply_to(message,"Yikes! No books for that genre. Maybe try another one?")
        return
    output = io.StringIO()
    csv_writer = csv.writer(output)
    csv_writer.writerow(['Title', 'Authors', 'Description', 'Preview Link'])

    for book in books:
        title = book['volumeInfo'].get('title', 'No title')
        authors = ', '.join(book['volumeInfo'].get('authors', ['Unknown author']))
        description = book['volumeInfo'].get('description', 'No description available')
        preview_link = book['volumeInfo'].get('previewLink', 'No preview available')
        csv_writer.writerow([title, authors, description, preview_link])

    output.seek(0)
    bot.send_document(message.chat.id, io.BytesIO(output.getvalue().encode()), caption="Boom! Here's a CSV with book deets. 📄📚")

@bot.message_handler(commands=['preview'])
def book_enquiry(message):
    msg = bot.reply_to(message, "What book's got you curious? Drop its name and I'll grab the preview link for ya!")
    bot.register_next_step_handler(msg, preview_link)

def preview_link(message):
    book_name = message.text
    results = search_books(book_name)
    books = results.get('items', [])
    
    if not books:
        bot.reply_to(message, "Oopsie! No preview link found. Double-check the book name and try again?")
        return
    
    preview_link = books[0]['volumeInfo'].get('previewLink', 'No preview available')
    if preview_link == 'No preview available':
        bot.reply_to(message, preview_link)
    else:
        bot.reply_to(message, f"Check this out: {preview_link}")



@bot.message_handler(commands=['list'])
def book_enquiry_list(message):
    msg = bot.reply_to(message, "Ready to boss up your reading list? Tell me the book name and I’ll sort you out. 😉")
    bot.register_next_step_handler(msg, process_book_list)

def process_book_list(message):
    book_name = message.text
    bot.reply_to(message,"For list management, just hit up /reading_list. I’ll handle the rest!")

@bot.message_handler(commands=['reading_list'])
def show_reading_list(message):
    markup = types.ReplyKeyboardMarkup(row_width=2)
    btn_add = types.KeyboardButton('Add a book')
    btn_delete = types.KeyboardButton('Delete a book')
    btn_view = types.KeyboardButton('View Reading List')
    markup.add(btn_add, btn_delete, btn_view)
    bot.reply_to(message, "What’s the move with your reading list? Choose an option below:", reply_markup=markup)

def create_or_load_docx():
    if os.path.exists('reading_list.docx'):
        return Document('reading_list.docx')
    else:
        doc = Document()
        doc.add_heading('Reading List', level=1)
        doc.save('reading_list.docx')
        return doc

def add_to_reading_list(title):
    doc = create_or_load_docx()
    doc.add_paragraph(f"📖 Title: {title}\n", style='Normal')
    doc.save('reading_list.docx')

def delete_from_reading_list(title):
    doc = Document('reading_list.docx')
    new_doc = Document()
    for para in doc.paragraphs:
        if title not in para.text:
            new_doc.add_paragraph(para.text)
    new_doc.save('reading_list.docx')

@bot.message_handler(func=lambda message: message.text == 'Add a book')
def add_book(message):
    msg = bot.reply_to(message, "Got a book in mind? Drop the title here.")
    bot.register_next_step_handler(msg, process_add_book)

def process_add_book(message):
    title = message.text.strip()  
    if not title:
        bot.reply_to(message,"Whoa there! You gotta give me a book title to add.")
        return
    add_to_reading_list(title)
    bot.reply_to(message, "Sweet! Your book is now on the reading list. 📚✨.\n Wanna checkout how far your reading list has come to /reading_list")

@bot.message_handler(func=lambda message: message.text == 'Delete a book')
def delete_book(message):
    msg = bot.reply_to(message, "Which book needs to go? Drop the title and it’s outta here.")
    bot.register_next_step_handler(msg, process_delete_book)

def process_delete_book(message):
    title = message.text.strip()
    delete_from_reading_list(title)
    bot.reply_to(message, "All done! That book’s been erased from your reading list. 🗑️\n Wanna checkout how far your reading list has come to /reading_list")

@bot.message_handler(func=lambda message: message.text == 'View Reading List')
def view_reading_list(message):
    with open('reading_list.docx', 'rb') as docx_file:
        bot.send_document(message.chat.id, docx_file)


bot.infinity_polling()
